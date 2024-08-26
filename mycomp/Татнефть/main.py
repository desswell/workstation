from pathlib import Path
import io
import base64
from PIL import Image
import docx
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.text.run import CT_R
from docx.table import Table
from docx.text.paragraph import Paragraph
import itertools
import re
from transformers import T5ForConditionalGeneration, T5Tokenizer
import torch
from natasha import Doc, Segmenter, MorphVocab, NewsEmbedding, NewsMorphTagger
import numpy as np
from datetime import datetime
import openpyxl
from openpyxl.styles import Font, Alignment, NamedStyle
import re

import lxml
from lxml import etree
from docx import Document
from docx.shared import Pt




class DocumentParser(object):
    """
    Класс первичного париснга документа docx
    :code_assign: service
    :code_type: Пользовательские функции
    :packages:
    from pathlib import Path
    import io
    import base64
    from PIL import Image
    import docx
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.text.run import CT_R
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    """

    def __init__(self, filename, image_as='base64', image_type='png', media_dir='media', **kwargs):
        """
        :param filename: filename
        :param image_as: image encoding
        :param image_type: image type
        :param media_dir: media directory
        """
        self.image_as = image_as
        self.media_dir = Path(media_dir)
        self.image_type = image_type
        self.document = docx.Document(filename)

    def get_element_text(self, element):
        """get all text of an element
        """
        try:
            children = element.xpath('.//w:t')  # not working for lxml.etree._Element
        except:
            children = element.iterchildren()
        return ''.join(c.text for c in children if c.text).strip()

    def blob_to_image(self, blob,
                      image_as='base64',
                      image_type='jpeg',
                      filename='image',
                      media_dir=Path('../../../../../Downloads/Telegram Desktop'),
                      ):
        """convet image blob data to image file or base64 string
        """
        image = Image.open(io.BytesIO(blob))

        if image_type.lower() in ('jpeg', 'jpg'):
            image_type = 'jpeg'
            image = image.convert('RGB')  # png => jpeg, smaller size
            filename = f'{filename}.jpg'
        else:
            filename = f'{filename}.png'

        if image_as == 'file':
            if not media_dir.exists():
                media_dir.mkdir(parents=True)
            image.save(media_dir.joinpath(filename))
            image = str(media_dir.joinpath(filename))
        else:
            buffered = io.BytesIO()
            image.save(buffered, image_type)
            prefix = f'data:image/{image_type};base64,'.encode()
            image = prefix + base64.b64encode(buffered.getvalue())
            image = image.decode()

        return image, filename

    def parse(self):
        for n, element in enumerate(self.document.element.body.iterchildren()):
            if isinstance(element, CT_P):
                yield from self.parse_paragraph(Paragraph(element, self.document))
            elif isinstance(element, CT_Tbl):
                yield self.parse_table(Table(element, self.document))

    def parse_paragraph(self, paragraph):
        """parse paragraph
        """
        if paragraph._element.xpath('.//a:graphic|.//w:hyperlink'):
            yield 'multipart', self._parse_child_paragraph(paragraph._element)
        else:
            text = self.get_element_text(paragraph._element)
            if text:
                yield 'paragraph', {'text': text, 'style_id': paragraph.style.style_id}

    def _parse_child_paragraph(self, element):
        """parse paragraph with graphic or hyperlink
        """
        data = []
        for child in element.iterchildren():

            if isinstance(child, CT_R) and child.xpath('.//a:graphic'):
                part = self._parse_graphic(child)
            elif isinstance(child, lxml.etree._Element):
                part = self._parse_hyperlink(child)
            else:
                part = self.get_element_text(child)
                if not part:
                    continue
            data.append(part)
        return data

    def _parse_graphic(self, element):
        """parse graphic, return image data
        """
        rid = element.xpath('.//a:blip/@*')[0]
        im = self.document.part.rels[rid]._target
        image, filename = self.blob_to_image(
            im.blob,
            image_as=self.image_as,
            image_type=self.image_type,
            filename=im.sha1,
            media_dir=self.media_dir)
        return {
            'type': self.image_as,
            'filename': filename,
            'image': image,
        }

    def _parse_hyperlink(self, element):
        """parse hyperlink, return text and href
        """
        for value in element.values():
            if value.startswith('rId'):
                href = self.document.part.rels[value]._target
                text = self.get_element_text(element)
                return {'text': text, 'href': href}

    def parse_table(self, table, strip=True):
        """parse table, return table data and merged_cells
        """
        data = [
            [cell.text.strip() if strip else cell.text for cell in row.cells]
            for row in table.rows
        ]

        merged_cells = {}
        for x, row in enumerate(table.rows):
            for y, cell in enumerate(row.cells):
                if cell._tc.vMerge or cell._tc.grid_span != 1:
                    tc = (cell._tc.top, cell._tc.bottom,
                          cell._tc.left, cell._tc.right)
                    merged_cells['_'.join(map(str, tc))] = cell.text

        return 'table', {'data': data, 'merged_cells': merged_cells}


def parse_doc(path):
    """
    Функция для первичного парсинга документа

    :code_assign: service
    :code_type: Пользовательские функции

    :imports: DocumentParser

    :packages:
    import itertools
    import re

    :param str path: путь до документа docx

    :returns: results_concat, stats
    :rtype: list, dict
    """
    stats = dict()
    stats['Заголовки'] = 0
    stats['Таблицы'] = 0
    stats['Параграфы'] = 0
    stats['Изображения'] = 0
    stats['Слова'] = 0

    # parse document from xml
    doc = DocumentParser(path)
    print('doc', doc)
    parsed = list(doc.parse())
    # get different types of conntent, re-format the data
    # tables = [(j, item) for (j, item) in enumerate(parsed) if item[0] == 'table']
    paragraphs = [(j, item) for (j, item) in enumerate(parsed) if item[0] == 'paragraph']
    headings = [(j, item) for (j, item) in paragraphs if
                item[1].get('style_id') in ['Heading' + str(i) for i in range(8)]]
    images = [im for im in [el for el in itertools.chain(*[item[1] for item in parsed if item[0] == 'multipart']) if el]
              if im.get('image')]

    stats['Изображения'] = len(images)

    # Если есть разметка, то мы разбираем документ начиная с первого заголовка
    # В противном случае, начинаем разбирать все извлеченные элементы
    if len(headings) == 0:
        start_position = 0
    else:
        start_position = headings[0][0]

    # переформатируем элементы документа
    # формат {'<категория_элемента>: <содержание (текст или таблица в виде листа листов)>'}
    results = []
    for item in parsed[start_position:]:
        if item[0] == 'table':
            results.append({'Таблица': item[1]['data']})
            stats['Таблицы'] += 1
        elif item[0] == 'paragraph':
            style = item[1].get('style_id')
            if style in ['Heading' + str(i) for i in range(8)]:
                results.append({'Заголовок ' + style[-1]: item[1]['text']})
                stats['Заголовки'] += 1
            else:
                results.append({'Текст': item[1]['text']})

    # соединяем разрозненные текстовые элементы в один параграф.
    # Параграфы разделяются между собой таблицами и заголовками.
    document = []
    cumulative_text = ''
    for item in results:
        if item.get('Текст'):
            cumulative_text = cumulative_text + item.get('Текст')
            if cumulative_text[-1] not in ',.:;!?-–':
                cumulative_text = cumulative_text + '. '
            else:
                cumulative_text = cumulative_text + ' '
        else:
            if len(cumulative_text) > 0:
                document.append({'Текст': cumulative_text})
                stats['Параграфы'] += 1
            document.append(item)
            cumulative_text = ''
    if cumulative_text != '':
        document.append({'Текст': cumulative_text})

    print(document)
    # сборка содержания
    # определяем все типы элементов, которые мы извлекли из документа
    headings_hierarchy = set([list(element.keys())[0] for element in document])
    if 'Текст' in headings_hierarchy:
        headings_hierarchy.remove('Текст')

    contents_tree = []
    for pos, item in enumerate(document):
        if any(key in headings_hierarchy for key in list(item.keys())):
            if list(item.keys())[0] == 'Таблица':
                # ищем заголовок для таблицы
                key = list(document[pos - 1].keys())[0]
                # Берём заголовок как последнее предложение в предшествующем парграфе, если оно содержит слово "Таблица"
                if key == 'Текст':
                    text = document[pos - 1].get(key)
                    sentences = re.split('[?.!]', text)
                    if re.search('Таблица', sentences[-2]):
                        contents_tree.append((pos, list(item.keys())[0], sentences[-2]))
                    else:
                        # в другом случае, название таблицы не передано
                        contents_tree.append((pos, list(item.keys())[0], 'Таблица - Безымянная таблица'))
                elif key in [zag for zag in ['Заголовок ' + str(i) for i in range(1, 7)]]:
                    text = document[pos - 1].get(key)
                    contents_tree.append((pos, list(item.keys())[0], 'Таблица - ' + text))
            else:
                contents_tree.append((pos, list(item.keys())[0], item.get(list(item.keys())[0])))

    stats['Содержание'] = list()
    counter = [0]
    for i, item in enumerate(contents_tree):
        name = item[2]
        name = name[name.find(next(filter(str.isalpha, name))):]
        if (item[1] == 'Заголовок 1') or (len(stats['Содержание']) == 0):
            counter = [counter[0] + 1]
            stats['Содержание'].append((i, item[1], name, str('.'.join([str(num) for num in counter]))))
        else:
            if item[1] > stats['Содержание'][-1][1]:
                counter.append(1)
            elif item[1] == stats['Содержание'][-1][1]:
                counter[-1] += 1
            else:
                counter = counter[:-1]
                counter[-1] += 1
            stats['Содержание'].append((i, item[1], name, '.'.join([str(num) for num in counter])))
    stats['Содержание'] = [(item[-1] + '.', item[2]) for item in stats['Содержание']]

    for element in document:
        if element.get('Текст'):
            stats['Слова'] += len(element.get('Текст').split())
        elif element.get('Таблица'):
            for row in element.get('Таблица'):
                for cell in row:
                    stats['Слова'] += len(cell.split())
        else:
            stats['Слова'] += sum([len(el.split()) for el in list(element.values())])

    return document, stats

def search_table(table_items: dict, keywords: list, title: str = None):
    """ ... Функция для поиска по словам-триггерам в таблице :code_assign: service :code_type: Пользовательские функции :imports: MystemUpdate :param dict table_items: обрабатываемый словарь таблицы :param list keywords: список слов-триггеров :param str table: название таблицы :returns: results_concat :rtype: list """
    stemmer = Mystem()
    # prepare the table
    table_items = table_items['Таблица']
    header = table_items[0]
    rows = list()
    # filter out if row has different number of elements
    for i in range(1, len(table_items)):
        if len(table_items[i]) == len(header):
            rows.append(table_items[i])
    output = []
    # search the table
    if title:
        triggered = set(list(filter(lambda x: x.isalnum(), stemmer.lemmatize(title)))).intersection(set(keywords))
        if len(triggered) > 0:
            output_column = list()
            for row in rows:
                output_column.append(', '.join(row))
            output.append((f'{";".join(output_column)}', triggered))
    else:
        for col_id, col_name in enumerate(header):
            triggered = set(list(filter(lambda x: x.isalnum(), stemmer.lemmatize(col_name)))).intersection(
                set(keywords))
            if len(triggered) > 0:
                output_column = []
                for row in rows:
                    output_column.append(rows[col_id])
                output.append((f'{col_name}: {";".join(output_column)}', triggered))
    return output

# def search_table(table: dict, keywords: list, title: str = None):
#     """
#     Функция для поиска по словам-триггерам в таблице

#     :code_assign: service
#     :code_type: Пользовательские функции

#     :imports: MystemUpdate

#     :param dict table: обрабатываемый словарь таблицы
#     :param list keywords: список слов-триггеров
#     :param str table: название таблицы

#     :returns: results_concat
#     :rtype: list
#     """

#     stemmer = Mystem()

#     # prepare the table
#     table_items = table['Таблица']
#     header = table_items[0]
#     rows = list()
#     # filter out if row has different number of elements
#     for i in range(1, len(table_items)):
#         if len(table_items[i]) == len(header):
#             rows.append(table_items[i])

#     output = []

#     # search the table 3 scenarios:
#     if title:
#         triggered = set(list(filter(lambda x: x.isalnum(), stemmer.lemmatize(title)))).intersection(set(keywords))
#         if len(triggered) > 0:
#             output_column = list()
#             for row in rows:
#                 output_column.append(', '.join(row))
#             output.append((f'{";".join(output_column)}', triggered))
#     else:
#         for col_id, col_name in enumerate(header):
#             triggered = set(list(filter(lambda x: x.isalnum(), stemmer.lemmatize(col_name)))).intersection(
#                 set(keywords))
#             if len(triggered) > 0:
#                 output_column = []
#                 for row in rows:
#                     output_column.append(rows[col_id])
#                 output.append((f'{col_name}: {";".join(output_column)}', triggered))
#     return output

def trigger_search(data: list):
    """ Функция для поиска по словам-триггерам по документу
    :code_assign: service :code_type: Пользовательские функции
    :imports: MystemUpdate, search_table
    :packages: from natasha import Doc,Segmenter
    :param list data: обрабатываемый документ в виде списка элемнтов-словарей
    :returns: matched_text
    :rtype: list """
    matched_text = list()

    def split_text_into_sentences(text):
        # Шаг 1: Замена "г." на временные метки
        text = re.sub(r'(\d{1,2}\.\d{1,2}\.\d{4})\sг\.\s(?=[А-Я])', r'\1__END_G__', text)
        text = re.sub(r'(\d{4})\sг\.\s(?=[А-Я])', r'\1__END_G__', text)
        text = re.sub(r'\bг\.', '__YEAR_G__', text)

        # Шаг 2: Добавление меток для границ предложений, включая __END_G__
        sentence_endings = re.compile(r'(?<!__YEAR_G__)\s*([.!?])\s+(?=[А-Я\d])|__END_G__')
        text = sentence_endings.sub(lambda m: f"{m.group(0).strip()}__SPLIT__", text)

        # Шаг 3: Восстановление точек в сокращениях
        text = text.replace('__YEAR_G__', 'г.')
        text = text.replace('__END_G__', ' г.')

        # Шаг 4: Разделение текста на предложения по меткам
        sentences = text.split('__SPLIT__')

        return [sentence.strip() for sentence in sentences]

    to_split = data[0]['Текст']
    res = split_text_into_sentences(to_split)
    print('ress', res)

    def filter_sentences_with_keywords(sentences):
        # Регулярное выражение для поиска дат, времен года, месяцев, годов и ключевых слов
        date_time_keywords_pattern = re.compile(
            r'(\d{1,2}\.\d{1,2}\.\d{4})|'  # дата в формате дд.мм.гггг
            r'(\d{1,2}\.\d{1,2}\.\d{2})|'  # дата в формате дд.мм.гг
            r'(\d{1,2}/\d{1,2}/\d{4})|'    # дата в формате дд/мм/гггг
            r'(\d{1,2}/\d{1,2}/\d{2})|'    # дата в формате дд/мм/гг
            r'(\d{1,2}-\d{1,2}-\d{2})|'    # дата в формате дд-мм-гггг
            r'(\d{1,2}-\d{1,2}-\d{4})|'    # дата в формате дд-мм-гггг
            r'\b(январ[ьяею]|феврал[ьяею]|март[аею]?|апрел[ьяею]|ма[ейя]|июн[ьяюе]?|июл[ьяюе]?|август[аею]?|'
            r'сентябр[ьяею]|октябр[ьяею]|ноябр[ьяею]|декабр[ьяею])\b|'
            r'\b(момент[ауеом]?|сейчас|сегодня|состояни[еяюем]|времен[еиемяю]|ситуаци[еяюейию]|период[аеуом]?|накануне|начал[оаиу]?|'
            r'кон[еуцаом]?|перспектив[аеуойию]|дн[еяю]|день|месяц[аеуом]?|год[ауеом]?|недел[яюеийямеи])\b',
            re.IGNORECASE
        )

        filtered_sentences = [sentence for sentence in sentences if date_time_keywords_pattern.search(sentence)]
        return filtered_sentences

    filtered_result = filter_sentences_with_keywords(res)
    print('sssssss', filtered_result)
    print(len(filtered_result))
    return filtered_result

# def trigger_search(data: list, keywords: list):
#     """
#     Функция для поиска по словам-триггерам по документу

#     :code_assign: service
#     :code_type: Пользовательские функции

#     :imports: MystemUpdate, search_table

#     :packages:
#     from natasha import Doc,Segmenter

#     :param list data: обрабатываемый документ в виде списка элемнтов-словарей
#     :param list keywords: список слов-триггеров

#     :returns: matched_text
#     :rtype: list
#     """

#     # aux functions
#     stemmer = Mystem()

#     # iteration
#     matched_text = list()
#     for i, element in enumerate(data):
#         # parse text
#         if element.get('Текст'):
#             text = element.get('Текст')
#             doc = Doc(text)
#             doc.segment(Segmenter())

#             for sent in doc.sents:
#                 words_list = list(filter(lambda x: x.isalnum(), stemmer.lemmatize(sent.text)))
#                 # if any(trigger in words_list for trigger in keywords):
#                 #     matched_text.append(sent.text)
#                 triggered = set(words_list).intersection(set(keywords))
#                 if len(triggered) > 0:
#                     matched_text.append((sent.text, triggered))

#         # прорабатываем таблицы
#         elif element.get('Таблица'):
#             title = None

#             # try get title
#             key = list(data[i - 1].keys())[0]
#             if key == 'Текст':
#                 text = data[i - 1].get(key)
#                 doc = Doc(text)
#                 doc.segment(Segmenter())
#                 if 'таблица' in list(filter(lambda x: x.isalnum(), stemmer.lemmatize(doc.sents[-1].text))):
#                     title = doc.sents[-1].text
#             elif key in [zag for zag in ['Заголовок ' + str(i) for i in range(1, 7)]]:
#                 text = data[i - 1].get(key)
#                 title = text

#             matched_text.extend(search_table(element, keywords, title))
#         # ничего не делаем с заголовками
#         else:
#             pass

#     return matched_text

# def summarize_texts(texts: list):
#     """ Функция для саммаризации предложений :code_assign: service :code_type: Пользовательские функции :packages: from transformers import T5ForConditionalGeneration, T5Tokenizer import torch :param list texts: отобранные тексты :returns: gui_dict, error :rtype: dict, str """
#     # prepare the model
#     MODEL_NAME = 't5-base'
#     model = T5ForConditionalGeneration.from_pretrained(MODEL_NAME)
#     tokenizer = T5Tokenizer.from_pretrained(MODEL_NAME, cache_dir="./")
#     model.eval()
#     output_texts = list()
#     for (text, trigger) in texts:
#         if len(text) < 15000:
#             input_text = text
#             input_tokens = tokenizer(input_text, return_tensors='pt', padding=True).to(model.device)
#             with torch.inference_mode():
#                 output_tokens = model.generate(**input_tokens, max_length=1000, num_beams=3, do_sample=False,
#                                               repetition_penalty=10.0)
#             output_texts.append((tokenizer.decode(output_tokens[0], skip_special_tokens=True), trigger))
#     return output_texts

# def summarize_texts(texts: list):
#     """
#     Функция для саммаризации предложений

#     :code_assign: service
#     :code_type: Пользовательские функции

#     :packages:
#     from transformers import T5ForConditionalGeneration, T5Tokenizer
#     import torch

#     :param list texts: отобранные тексты

#     :returns: gui_dict, error
#     :rtype: dict, str
#     """

#     MODEL_NAME = 't5-base'

#     model = T5ForConditionalGeneration.from_pretrained(MODEL_NAME)
#     tokenizer = T5Tokenizer.from_pretrained(MODEL_NAME, cache_dir="./")
#     model.eval()

#     output_texts = list()
#     for (text, trigger) in texts:
#         if len(text) < 15000:
#             input_text = '[{}] '.format(100) + text
#             input_tokens = tokenizer(input_text, return_tensors='pt', padding=True).to(model.device)
#             with torch.inference_mode():
#                 output_tokens = model.generate(**input_tokens,
#                                                max_length=1000, num_beams=3,
#                                                do_sample=False, repetition_penalty=10.0)
#             output_texts.append((tokenizer.decode(output_tokens[0], skip_special_tokens=True), trigger))

#     return output_texts


def get_document_stats(document: list):
    """
    Функция для статистического анализа текста

    :code_assign: service
    :code_type: Пользовательские функции

    :packages:
    import re
    from natasha import Doc, Segmenter, MorphVocab, NewsEmbedding, NewsMorphTagger
    import numpy as np

    :param list document: лист с элементами документа после парсинга

    :returns: stats
    :rtype: dict
    """
    plain_text = list()
    for item in document:
        if item.get('Таблица'):
            table = item.get('Таблица')
            plain_text.append('. '.join(('; '.join(cell for cell in row)) for row in table))
        else:
            plain_text.append('. '.join(list(item.values())))
    plain_text = '. '.join(plain_text)

    with open('../../../Downloads/Telegram Desktop/stopwords_ru.txt') as f:
        russian_stopwords = f.read().splitlines()

    segmenter = Segmenter()
    morph_vocab = MorphVocab()
    emb = NewsEmbedding()
    morph_tagger = NewsMorphTagger(emb)

    text = Doc(plain_text)
    text.segment(segmenter)
    text.tag_morph(morph_tagger)

    for token in text.tokens:
        token.lemmatize(morph_vocab)

    # рассчёт статистики по всем словам
    nouns = [token.lemma for token in text.tokens if token.pos in ['PROPN', 'NOUN']]
    adjectives = [token.lemma for token in text.tokens if token.pos in ['ADJ']]

    unique_nouns, counts_nouns = np.unique(nouns, return_counts=True)
    sorted_nouns = list(
        filter(lambda x: x[1] not in russian_stopwords, sorted(zip(counts_nouns, unique_nouns), reverse=True)))

    unique_adj, counts_adj = np.unique(adjectives, return_counts=True)
    sorted_adj = list(
        filter(lambda x: x[1] not in russian_stopwords, sorted(zip(counts_adj, unique_adj), reverse=True)))

    # рассчёт статистики по терминам
    # Поиск терминов с помощью регулярных выражений
    terms = [x.replace('\xa0', ' ') for x in re.findall(r'\«(.*?)\»', plain_text)]
    # Добавление английских терминов из текста
    english_terms = re.findall(r'\b[A-Z][A-Za-z0-9]+\b', plain_text)
    terms.extend(english_terms)
    unique_terms, counts_terms = np.unique(terms, return_counts=True)
    sorted_terms = sorted(zip(counts_terms, unique_terms), reverse=True)

    stats = dict()
    stats['Существительные'] = sorted_nouns[:20]
    stats['Прилагательные'] = sorted_adj[:20]
    stats['Термины'] = sorted_terms[:20]

    return stats


def save_to_docx(result_path, trigger_words, output: str = None):
    """
    Функция для компановки отчёта в docx файле

    :code_assign: service
    :code_type: Пользовательские функции

    :packages:
    from docx import Document
    :param str result_path: результат
    :param list trigger_words: список слов триггеров с номером раздела, к которому слово относится
    :param str output: название выходного файла
    """
    doc = Document()

    # Добавление предложений в документ
    for sentence in trigger_words:
        doc.add_paragraph(sentence)

    # Сохранение документа
    doc.save(output)

# пользовательская

def document_analysis(
        start: bool,
        document_path: str,
        output_name: str = None
):
    """
    Функция для анализа текста
    :code_assign: users
    :code_type: Пользовательские функции
    :imports: init_gui_dict, parse_doc, trigger_search, MystemUpdate, summarize_texts, save_to_docx, get_document_stats
    :packages:
    import re
    from natasha import Doc, Segmenter
    :param_block bool start: start
    :param str document_path: путь до документа docx
    :param str output_name: название выходного файла xlsx
    :returns: gui_dict, error
    :rtype: dict, str
    :semrtype: ,
    """
    gui_dict = init_gui_dict()
    error = ''
    document, stats = parse_doc(document_path)
    stats.update(get_document_stats(document))
    output = trigger_search(document)

    save_to_docx("", output, "result.docx")


    return gui_dict, error


file_path = 'тестовый_документ.docx'

document, stats = parse_doc(file_path)

print(stats)

stats.update(get_document_stats(document))

trigger_words = trigger_search(document)

save_to_docx("", trigger_words, "result.docx")
