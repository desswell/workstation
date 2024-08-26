from pathlib import Path
import io
import base64
from PIL import Image
import docx
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.text.run import CT_R
from docx.table import Table
from docx.text.paragraph import Paragraph
import itertools
import re
from transformers import T5ForConditionalGeneration, T5Tokenizer
import torch
from natasha import Doc, Segmenter, MorphVocab, NewsEmbedding, NewsMorphTagger
import numpy as np
from datetime import datetime
import openpyxl
from openpyxl.styles import Font, Alignment, NamedStyle
import re

import lxml
from lxml import etree
from docx import Document
from docx.shared import Pt




class DocumentParser(object):
    """
    Класс первичного париснга документа docx
    :code_assign: service
    :code_type: Пользовательские функции
    :packages:
    from pathlib import Path
    import io
    import base64
    from PIL import Image
    import docx
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.text.run import CT_R
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    """

    def __init__(self, filename, image_as='base64', image_type='png', media_dir='media', **kwargs):
        """
        :param filename: filename
        :param image_as: image encoding
        :param image_type: image type
        :param media_dir: media directory
        """
        self.image_as = image_as
        self.media_dir = Path(media_dir)
        self.image_type = image_type
        self.document = docx.Document(filename)

    def get_element_text(self, element):
        """get all text of an element
        """
        try:
            children = element.xpath('.//w:t')  # not working for lxml.etree._Element
        except:
            children = element.iterchildren()
        return ''.join(c.text for c in children if c.text).strip()

    def blob_to_image(self, blob,
                      image_as='base64',
                      image_type='jpeg',
                      filename='image',
                      media_dir=Path('../../../../../Downloads/Telegram Desktop'),
                      ):
        """convet image blob data to image file or base64 string
        """
        image = Image.open(io.BytesIO(blob))

        if image_type.lower() in ('jpeg', 'jpg'):
            image_type = 'jpeg'
            image = image.convert('RGB')  # png => jpeg, smaller size
            filename = f'{filename}.jpg'
        else:
            filename = f'{filename}.png'

        if image_as == 'file':
            if not media_dir.exists():
                media_dir.mkdir(parents=True)
            image.save(media_dir.joinpath(filename))
            image = str(media_dir.joinpath(filename))
        else:
            buffered = io.BytesIO()
            image.save(buffered, image_type)
            prefix = f'data:image/{image_type};base64,'.encode()
            image = prefix + base64.b64encode(buffered.getvalue())
            image = image.decode()

        return image, filename

    def parse(self):
        for n, element in enumerate(self.document.element.body.iterchildren()):
            if isinstance(element, CT_P):
                yield from self.parse_paragraph(Paragraph(element, self.document))
            elif isinstance(element, CT_Tbl):
                yield self.parse_table(Table(element, self.document))

    def parse_paragraph(self, paragraph):
        """parse paragraph
        """
        if paragraph._element.xpath('.//a:graphic|.//w:hyperlink'):
            yield 'multipart', self._parse_child_paragraph(paragraph._element)
        else:
            text = self.get_element_text(paragraph._element)
            if text:
                yield 'paragraph', {'text': text, 'style_id': paragraph.style.style_id}

    def _parse_child_paragraph(self, element):
        """parse paragraph with graphic or hyperlink
        """
        data = []
        for child in element.iterchildren():

            if isinstance(child, CT_R) and child.xpath('.//a:graphic'):
                part = self._parse_graphic(child)
            elif isinstance(child, lxml.etree._Element):
                part = self._parse_hyperlink(child)
            else:
                part = self.get_element_text(child)
                if not part:
                    continue
            data.append(part)
        return data

    def _parse_graphic(self, element):
        """parse graphic, return image data
        """
        rid = element.xpath('.//a:blip/@*')[0]
        im = self.document.part.rels[rid]._target
        image, filename = self.blob_to_image(
            im.blob,
            image_as=self.image_as,
            image_type=self.image_type,
            filename=im.sha1,
            media_dir=self.media_dir)
        return {
            'type': self.image_as,
            'filename': filename,
            'image': image,
        }

    def _parse_hyperlink(self, element):
        """parse hyperlink, return text and href
        """
        for value in element.values():
            if value.startswith('rId'):
                href = self.document.part.rels[value]._target
                text = self.get_element_text(element)
                return {'text': text, 'href': href}

    def parse_table(self, table, strip=True):
        """parse table, return table data and merged_cells
        """
        data = [
            [cell.text.strip() if strip else cell.text for cell in row.cells]
            for row in table.rows
        ]

        merged_cells = {}
        for x, row in enumerate(table.rows):
            for y, cell in enumerate(row.cells):
                if cell._tc.vMerge or cell._tc.grid_span != 1:
                    tc = (cell._tc.top, cell._tc.bottom,
                          cell._tc.left, cell._tc.right)
                    merged_cells['_'.join(map(str, tc))] = cell.text

        return 'table', {'data': data, 'merged_cells': merged_cells}


def parse_doc(path):
    """
    Функция для первичного парсинга документа

    :code_assign: service
    :code_type: Пользовательские функции

    :imports: DocumentParser

    :packages:
    import itertools
    import re

    :param str path: путь до документа docx

    :returns: results_concat, stats
    :rtype: list, dict
    """
    stats = dict()
    stats['Заголовки'] = 0
    stats['Таблицы'] = 0
    stats['Параграфы'] = 0
    stats['Изображения'] = 0
    stats['Слова'] = 0

    # parse document from xml
    doc = DocumentParser(path)
    print('doc', doc)
    parsed = list(doc.parse())
    # get different types of conntent, re-format the data
    # tables = [(j, item) for (j, item) in enumerate(parsed) if item[0] == 'table']
    paragraphs = [(j, item) for (j, item) in enumerate(parsed) if item[0] == 'paragraph']
    headings = [(j, item) for (j, item) in paragraphs if
                item[1].get('style_id') in ['Heading' + str(i) for i in range(8)]]
    images = [im for im in [el for el in itertools.chain(*[item[1] for item in parsed if item[0] == 'multipart']) if el]
              if im.get('image')]

    stats['Изображения'] = len(images)

    # Если есть разметка, то мы разбираем документ начиная с первого заголовка
    # В противном случае, начинаем разбирать все извлеченные элементы
    if len(headings) == 0:
        start_position = 0
    else:
        start_position = headings[0][0]

    # переформатируем элементы документа
    # формат {'<категория_элемента>: <содержание (текст или таблица в виде листа листов)>'}
    results = []
    for item in parsed[start_position:]:
        if item[0] == 'table':
            results.append({'Таблица': item[1]['data']})
            stats['Таблицы'] += 1
        elif item[0] == 'paragraph':
            style = item[1].get('style_id')
            if style in ['Heading' + str(i) for i in range(8)]:
                results.append({'Заголовок ' + style[-1]: item[1]['text']})
                stats['Заголовки'] += 1
            else:
                results.append({'Текст': item[1]['text']})

    # соединяем разрозненные текстовые элементы в один параграф.
    # Параграфы разделяются между собой таблицами и заголовками.
    document = []
    cumulative_text = ''
    for item in results:
        if item.get('Текст'):
            cumulative_text = cumulative_text + item.get('Текст')
            if cumulative_text[-1] not in ',.:;!?-–':
                cumulative_text = cumulative_text + '. '
            else:
                cumulative_text = cumulative_text + ' '
        else:
            if len(cumulative_text) > 0:
                document.append({'Текст': cumulative_text})
                stats['Параграфы'] += 1
            document.append(item)
            cumulative_text = ''
    if cumulative_text != '':
        document.append({'Текст': cumulative_text})

    print(document)
    # сборка содержания
    # определяем все типы элементов, которые мы извлекли из документа
    headings_hierarchy = set([list(element.keys())[0] for element in document])
    if 'Текст' in headings_hierarchy:
        headings_hierarchy.remove('Текст')

    contents_tree = []
    for pos, item in enumerate(document):
        if any(key in headings_hierarchy for key in list(item.keys())):
            if list(item.keys())[0] == 'Таблица':
                # ищем заголовок для таблицы
                key = list(document[pos - 1].keys())[0]
                # Берём заголовок как последнее предложение в предшествующем парграфе, если оно содержит слово "Таблица"
                if key == 'Текст':
                    text = document[pos - 1].get(key)
                    sentences = re.split('[?.!]', text)
                    if re.search('Таблица', sentences[-2]):
                        contents_tree.append((pos, list(item.keys())[0], sentences[-2]))
                    else:
                        # в другом случае, название таблицы не передано
                        contents_tree.append((pos, list(item.keys())[0], 'Таблица - Безымянная таблица'))
                elif key in [zag for zag in ['Заголовок ' + str(i) for i in range(1, 7)]]:
                    text = document[pos - 1].get(key)
                    contents_tree.append((pos, list(item.keys())[0], 'Таблица - ' + text))
            else:
                contents_tree.append((pos, list(item.keys())[0], item.get(list(item.keys())[0])))

    stats['Содержание'] = list()
    counter = [0]
    for i, item in enumerate(contents_tree):
        name = item[2]
        name = name[name.find(next(filter(str.isalpha, name))):]
        if (item[1] == 'Заголовок 1') or (len(stats['Содержание']) == 0):
            counter = [counter[0] + 1]
            stats['Содержание'].append((i, item[1], name, str('.'.join([str(num) for num in counter]))))
        else:
            if item[1] > stats['Содержание'][-1][1]:
                counter.append(1)
            elif item[1] == stats['Содержание'][-1][1]:
                counter[-1] += 1
            else:
                counter = counter[:-1]
                counter[-1] += 1
            stats['Содержание'].append((i, item[1], name, '.'.join([str(num) for num in counter])))
    stats['Содержание'] = [(item[-1] + '.', item[2]) for item in stats['Содержание']]

    for element in document:
        if element.get('Текст'):
            stats['Слова'] += len(element.get('Текст').split())
        elif element.get('Таблица'):
            for row in element.get('Таблица'):
                for cell in row:
                    stats['Слова'] += len(cell.split())
        else:
            stats['Слова'] += sum([len(el.split()) for el in list(element.values())])

    return document, stats

def search_table(table_items: dict, keywords: list, title: str = None):
    """ ... Функция для поиска по словам-триггерам в таблице :code_assign: service :code_type: Пользовательские функции :imports: MystemUpdate :param dict table_items: обрабатываемый словарь таблицы :param list keywords: список слов-триггеров :param str table: название таблицы :returns: results_concat :rtype: list """
    stemmer = Mystem()
    # prepare the table
    table_items = table_items['Таблица']
    header = table_items[0]
    rows = list()
    # filter out if row has different number of elements
    for i in range(1, len(table_items)):
        if len(table_items[i]) == len(header):
            rows.append(table_items[i])
    output = []
    # search the table
    if title:
        triggered = set(list(filter(lambda x: x.isalnum(), stemmer.lemmatize(title)))).intersection(set(keywords))
        if len(triggered) > 0:
            output_column = list()
            for row in rows:
                output_column.append(', '.join(row))
            output.append((f'{";".join(output_column)}', triggered))
    else:
        for col_id, col_name in enumerate(header):
            triggered = set(list(filter(lambda x: x.isalnum(), stemmer.lemmatize(col_name)))).intersection(
                set(keywords))
            if len(triggered) > 0:
                output_column = []
                for row in rows:
                    output_column.append(rows[col_id])
                output.append((f'{col_name}: {";".join(output_column)}', triggered))
    return output

def trigger_search(data: list):
    """ Функция для поиска по словам-триггерам по документу
    :code_assign: service :code_type: Пользовательские функции
    :imports: MystemUpdate, search_table
    :packages: from natasha import Doc,Segmenter
    :param list data: обрабатываемый документ в виде списка элемнтов-словарей
    :returns: matched_text
    :rtype: list """
    matched_text = list()

    def split_text_into_sentences(text):
        # Шаг 1: Замена "г." на временные метки
        text = re.sub(r'(\d{1,2}\.\d{1,2}\.\d{4})\sг\.\s(?=[А-Я])', r'\1__END_G__', text)
        text = re.sub(r'(\d{4})\sг\.\s(?=[А-Я])', r'\1__END_G__', text)
        text = re.sub(r'\bг\.', '__YEAR_G__', text)

        # Шаг 2: Добавление меток для границ предложений, включая __END_G__
        sentence_endings = re.compile(r'(?<!__YEAR_G__)\s*([.!?])\s+(?=[А-Я\d])|__END_G__')
        text = sentence_endings.sub(lambda m: f"{m.group(0).strip()}__SPLIT__", text)

        # Шаг 3: Восстановление точек в сокращениях
        text = text.replace('__YEAR_G__', 'г.')
        text = text.replace('__END_G__', ' г.')

        # Шаг 4: Разделение текста на предложения по меткам
        sentences = text.split('__SPLIT__')

        return [sentence.strip() for sentence in sentences]

    to_split = data[0]['Текст']
    res = split_text_into_sentences(to_split)
    print('ress', res)

    def filter_sentences_with_keywords(sentences):
        # Регулярное выражение для поиска дат, времен года, месяцев, годов и ключевых слов
        date_time_keywords_pattern = re.compile(
            r'(\d{1,2}\.\d{1,2}\.\d{4})|'  # дата в формате дд.мм.гггг
            r'(\d{1,2}\.\d{1,2}\.\d{2})|'  # дата в формате дд.мм.гг
            r'(\d{1,2}/\d{1,2}/\d{4})|'    # дата в формате дд/мм/гггг
            r'(\d{1,2}/\d{1,2}/\d{2})|'    # дата в формате дд/мм/гг
            r'(\d{1,2}-\d{1,2}-\d{2})|'    # дата в формате дд-мм-гггг
            r'(\d{1,2}-\d{1,2}-\d{4})|'    # дата в формате дд-мм-гггг
            r'\b(январ[ьяею]|феврал[ьяею]|март[аею]?|апрел[ьяею]|ма[ейя]|июн[ьяюе]?|июл[ьяюе]?|август[аею]?|'
            r'сентябр[ьяею]|октябр[ьяею]|ноябр[ьяею]|декабр[ьяею])\b|'
            r'\b(момент[ауеом]?|сейчас|дней|настоящее время|сегодня|состояни[еяюем]|времен[еиемяю]|ситуаци[еяюейию]|вчера|позавчера|завтра|после завтра|период[аеуом]?|накануне|начал[оаиу]?|'
            r'кон[еуцаом]?|перспектив[аеуойию]|дн[еяю]|день|месяцев|месяц[аеуом]?|год[ауеом]?|недел[яюеийямеи])\b',
            re.IGNORECASE
        )

        filtered_sentences = [sentence for sentence in sentences if date_time_keywords_pattern.search(sentence)]
        return filtered_sentences

    filtered_result = filter_sentences_with_keywords(res)
    print('sssssss', filtered_result)
    print(len(filtered_result))
    return filtered_result



def get_document_stats(document: list):
    """
    Функция для статистического анализа текста

    :code_assign: service
    :code_type: Пользовательские функции

    :packages:
    import re
    from natasha import Doc, Segmenter, MorphVocab, NewsEmbedding, NewsMorphTagger
    import numpy as np

    :param list document: лист с элементами документа после парсинга

    :returns: stats
    :rtype: dict
    """
    plain_text = list()
    for item in document:
        if item.get('Таблица'):
            table = item.get('Таблица')
            plain_text.append('. '.join(('; '.join(cell for cell in row)) for row in table))
        else:
            plain_text.append('. '.join(list(item.values())))
    plain_text = '. '.join(plain_text)

    with open('stopwords_ru.txt') as f:
        russian_stopwords = f.read().splitlines()

    segmenter = Segmenter()
    morph_vocab = MorphVocab()
    emb = NewsEmbedding()
    morph_tagger = NewsMorphTagger(emb)

    text = Doc(plain_text)
    text.segment(segmenter)
    text.tag_morph(morph_tagger)

    for token in text.tokens:
        token.lemmatize(morph_vocab)

    # рассчёт статистики по всем словам
    nouns = [token.lemma for token in text.tokens if token.pos in ['PROPN', 'NOUN']]
    adjectives = [token.lemma for token in text.tokens if token.pos in ['ADJ']]

    unique_nouns, counts_nouns = np.unique(nouns, return_counts=True)
    sorted_nouns = list(
        filter(lambda x: x[1] not in russian_stopwords, sorted(zip(counts_nouns, unique_nouns), reverse=True)))

    unique_adj, counts_adj = np.unique(adjectives, return_counts=True)
    sorted_adj = list(
        filter(lambda x: x[1] not in russian_stopwords, sorted(zip(counts_adj, unique_adj), reverse=True)))

    # рассчёт статистики по терминам
    # Поиск терминов с помощью регулярных выражений
    terms = [x.replace('\xa0', ' ') for x in re.findall(r'\«(.*?)\»', plain_text)]
    # Добавление английских терминов из текста
    english_terms = re.findall(r'\b[A-Z][A-Za-z0-9]+\b', plain_text)
    terms.extend(english_terms)
    unique_terms, counts_terms = np.unique(terms, return_counts=True)
    sorted_terms = sorted(zip(counts_terms, unique_terms), reverse=True)

    stats = dict()
    stats['Существительные'] = sorted_nouns[:20]
    stats['Прилагательные'] = sorted_adj[:20]
    stats['Термины'] = sorted_terms[:20]

    return stats


def save_to_docx(result_path, trigger_words, output: str = None):
    """
    Функция для компановки отчёта в docx файле

    :code_assign: service
    :code_type: Пользовательские функции

    :packages:
    from docx import Document
    :param str result_path: результат
    :param list trigger_words: список слов триггеров с номером раздела, к которому слово относится
    :param str output: название выходного файла
    """
    doc = Document()

    # Добавление предложений в документ
    for sentence in trigger_words:
        doc.add_paragraph(sentence)

    # Сохранение документа
    doc.save(output)

# пользовательская

def document_analysis(
        start: bool,
        document_path: str,
        output_name: str = None
):
    """
    Функция для анализа текста
    :code_assign: users
    :code_type: Пользовательские функции
    :imports: init_gui_dict, parse_doc, trigger_search, MystemUpdate, summarize_texts, save_to_docx, get_document_stats
    :packages:
    import re
    from natasha import Doc, Segmenter
    :param_block bool start: start
    :param str document_path: путь до документа docx
    :param str output_name: название выходного файла xlsx
    :returns: gui_dict, error
    :rtype: dict, str
    :semrtype: ,
    """
    gui_dict = init_gui_dict()
    error = ''
    stemmer = MystemUpdate()

    # hard-coded слова
    keywords1 = ['контракт', 'соглашение', 'договоренность',
                 'исполнитель', 'эксперт', 'заказчик', 'клиент']
    keywords2 = ['срок', 'период', 'интервал', 'дата', 'день', 'месяц', 'год', 'календарь', ]
    keywords3 = ['цена', 'стоимость', 'плата', 'тариф', 'рубли', 'деньги', 'валюта']
    keywords4 = ['инструменты', 'БД', 'СУБД', 'сервер']
    keywords5 = ['график', 'фазы']
    keywords6 = ['требования', 'предписания', 'стандарты', 'рекомендации', 'параметры', 'спецификации', 'приоритеты']
    keywords7 = ['.01', '.02', '.03', '.04', '.05', '.06', '.07', '.08', '.09', '.10', '.11', '.12', 'март',
                 'апрель', 'май', 'июнь', 'июль', 'август', 'сентябрь', 'октябрь', 'ноябрь', 'декабрь', 'январь',
                 'февраль', 'день', 'месяц', 'год',
                 'момент', 'сейчас', 'сегодня', 'состояние', 'время', 'ситуация', 'период', 'накануне', 'начало',
                 'конец', 'перспектива']

    keyword_to_section = dict()
    for num, section in enumerate([keywords1, keywords2, keywords3, keywords4, keywords5, keywords6, keywords7]):
        for word in section:
            keyword_to_section[word] = num + 1
    document, stats = parse_doc(document_path)
    stats.update(get_document_stats(document))
    output = trigger_search(document)

    save_to_docx("", output, "result.docx")


    return gui_dict, error


file_path = 'тестовый_документ.docx'

document, stats = parse_doc(file_path)

print(stats)

#запуск функции save_to_excel
from pymystem3 import Mystem
keywords1 = ['контракт', 'соглашение', 'договоренность',
             'исполнитель', 'эксперт', 'заказчик', 'клиент']
keywords2 = ['срок', 'период', 'интервал', 'дата', 'день', 'месяц', 'год','календарь', ]
keywords3 = ['цена', 'стоимость', 'плата', 'тариф', 'рубли', 'деньги', 'валюта']
keywords4 = ['инструменты', 'БД', 'СУБД', 'сервер']
keywords5 = ['график', 'фазы']
keywords6 = ['требования', 'предписания', 'стандарты', 'рекомендации', 'параметры', 'спецификации', 'приоритеты']
keywords7 = ['.01','.02','.03','.04','.05','.06','.07','.08','.09','.10','.11','.12', 'март',
             'апрель', 'май', 'июнь', 'июль', 'август', 'сентябрь', 'октябрь', 'ноябрь', 'декабрь', 'январь', 'февраль', 'день', 'месяц', 'год',
             'момент', 'сейчас',  'сегодня', 'состояние', 'время', 'ситуация', 'период', 'накануне', 'начало', 'конец', 'перспектива']
keyword_to_section = dict()

# with open('stopwords_ru.txt', 'r') as file:ф
#     keywords = [line.strip() for line in file.readlines()]

# for num, section in enumerate([keywords]):
#     for word in section:
#         keyword_to_section[word] = num + 1

stemmer = Mystem()
keyword_to_section = dict()
for num, section in enumerate([keywords1, keywords2, keywords3, keywords4, keywords5, keywords6, keywords7]):
    for word in section:
        keyword_to_section[word] = num + 1


lemma_to_word = {lemma: orig for lemma, orig in zip(list(
    filter(lambda x: x.isalnum(), stemmer.lemmatize(' '.join(list(keyword_to_section.keys()))))),
    list(keyword_to_section.keys()))}
keywords = list(lemma_to_word.keys())

import nltk
from nltk.corpus import stopwords
# Загрузка стоп-слов для русского языка
nltk.download('stopwords')
stopwords_ru = set(stopwords.words('russian'))
# Путь к файлу stopwords_ru.txt
file_path = 'stopwords_ru.txt'
# Запись слов в файл
with open(file_path, 'w', encoding='utf-8') as file:
    file.write('\n'.join(stopwords_ru))

stats.update(get_document_stats(document))

trigger_words = trigger_search(document)

save_to_docx("", trigger_words, "result.docx")



import docx
import re

def convert_month_word_to_number(date_str):
    month_mapping = {
        'января': '01', 'февраля': '02', 'марта': '03', 'апреля': '04',
        'мая': '05', 'июня': '06', 'июля': '07', 'августа': '08',
        'сентября': '09', 'октября': '10', 'ноября': '11', 'декабря': '12'
    }

    for month_word, month_number in month_mapping.items():
        if month_word in date_str:
            day_month_match = re.search(r'\b(\d{1,2})\s' + month_word, date_str)
            if day_month_match:
                day = day_month_match.group(1).zfill(2)  # Ensure 2-digit day format
                month_number = str(int(month_number)).zfill(2)  # Ensure 2-digit month format
                if 'года' in date_str:
                    return re.sub(r'\b\d{1,2}\s' + month_word + r'\sгода', day + '.' + month_number + '.', date_str)
                else:
                    return re.sub(r'\b\d{1,2}\s' + month_word + r'\b', day + '.' + month_number + '.2024', date_str)
    return date_str

def process_today_date_in_text(paragraph, user_date):
    word_to_num = {
        'одного': 1, 'один': 1,
        'двух': 2, 'два': 2,
        'трех': 3, 'три': 3,
        'четырех': 4, 'четыре': 4,
        'пяти': 5, 'пять': 5,
        'шести': 6, 'шесть': 6,
        'семи': 7, 'семь': 7,
        'восьми': 8, 'восемь': 8,
        'девяти': 9, 'девять': 9,
        'десяти': 10, 'десять': 10
    }

    # Словарь для преобразования названий месяцев в числовые значения
    months = {
        'январь': 1, 'января': 1, 'январе': 1,
        'февраль': 2, 'февраля': 2, 'феврале': 2,
        'март': 3, 'марта': 3, 'марте': 3,
        'апрель': 4, 'апреля': 4,'апреле': 4,
        'май': 5, 'мая': 5, 'мае': 5,
        'июнь': 6, 'июня': 6, 'июне': 6,
        'июль': 7, 'июля': 7, 'июле': 7,
        'август': 8, 'августа': 8, 'августе': 8,
        'сентябрь': 9, 'сентября': 9, 'сентябре': 9,
        'октябрь': 10, 'октября': 10, 'октябре': 10,
        'ноябрь': 11, 'ноября': 11, 'ноябре': 11,
        'декабрь': 12, 'декабря': 12, 'декабре': 12
    }
    seasons = {
        'весна': 3,
        'лето': 6,
        'осень': 9,
        'зима': 12,
        'весной': 3,
        'летом': 6,
        'осенью': 9,
        'зимой': 12
    }
    patterns = [
        re.compile(r'(На\sсегодняшний\sдень|настоящее\sвремя|сегодня|сейчас)', re.IGNORECASE),
        re.compile(
            r'\bтекущ(?:ее|его|ему|ем|ая|ую|ей|ую|ие|их|им|ими|ее|ей|ие|ими)?\sсостояни(?:е|я|ю|ем|и|ям|ями|ях)?\b',
            re.IGNORECASE),
        re.compile(
            r'(\d{1,2}) (январ[ьяею]|феврал(?:ь|ю|е|я)|март|март[еа]|апрел(?:ь|ю|е|я)|ма[йея]|июн[ьея]|июл[ьея]|август|август[еa]|сентябр[ьея]|октябр(?:ь|ю|е|я)|ноябр[ьея]|декабр[ьея]) (\d{4})',
            re.IGNORECASE),
        re.compile(
            r'(\d{1,2})[./](\d{1,2})\s(январ[ьяею]|феврал(?:ь|ю|е|я)|март|март[еа]|апрел(?:ь|ю|е|я)|ма[йея]|июн[ьея]|июл[ьея]|август|август[еa]|сентябр[ьея]|октябр(?:ь|ю|е|я)|ноябр[ьея]|декабр[ьея])',
            re.IGNORECASE),
        re.compile(r'(\d{2}).(\d{2}).(\d{4})', re.IGNORECASE),  # DD-MM-YYYY
        re.compile(r'в течение ([а-я]+) (месяцев|год[ау]|лет|месяца|годов)', re.IGNORECASE),
        re.compile(r'с (\w+) по (\w+) (\d{4})', re.IGNORECASE),
        re.compile(r'(\d{2})/(\d{2})/(\d{2})', re.IGNORECASE),  # DD/MM/YY
        re.compile(r'(весна|весной|лето|летом|осень|осенью|зима|зимой) (\d{4})', re.IGNORECASE), # (Время года) (год)
        re.compile(r'(\d{1,2})(?:го)?\s(январ(?:е|ю|я)|феврал(?:ь|ю|е|я)|март|март[еа]|апрел(?:ь|ю|е|я)|ма[йея]|июн[ьея]|июл[ьея]|август|август[еa]|сентябр[ьея]|октябр(?:ь|ю|е|я)|ноябр[ьея]|декабр[ьея])', re.IGNORECASE),
        re.compile(r'(январ(?:е|ю)|феврал(?:ь|ю|е)|март|марте|апрел(?:ь|ю|е)|ма[йе]|июн[ье]|июл[ье]|август|августе|сентябр[ье]|октябр(?:ь|ю|е)|ноябр[ье]|декабр[ье]) (\d{4})',
                   re.IGNORECASE),
        re.compile(r'(\d{2})-(\d{2})-(\d{2})', re.IGNORECASE),  # DD-MM-YY
        re.compile(r'(\w+),?\sпрошлого года', re.IGNORECASE),  # (месяц),? прошлого года
        re.compile(r'(\w+),?\sследующего года', re.IGNORECASE),  # (месяц),? прошлого года
        re.compile(r'на прошлой неделе,?\s(\d{4})', re.IGNORECASE),
        re.compile(r'(\d{1,2})го\s(\w+)\s(\d{4})', re.IGNORECASE),
        re.compile(r'(\d{1,2})(?:го)?,?\s(\w+)\s(этого года|предыдущего года|следующего года)', re.IGNORECASE),
        re.compile(r'(весной|летом|осенью|зимой)\s(этого года|следующего года|предыдущего года)', re.IGNORECASE),
        re.compile(r'(\d{2})\.(\d{2})\.(\d{2})', re.IGNORECASE),
        re.compile(r'(\d{4})\sгод(?:у)?', re.IGNORECASE),
        re.compile(
            r'(вчера|позавчера|завтра|послезавтра)',
            re.IGNORECASE),
        re.compile(r'через\s(\d+)?\s?(месяц|месяца|месяцев|год|года|лет)', re.IGNORECASE),
        re.compile(r'через\s(день|\d+\sдней|\d+\sдня|несколько\sдней)', re.IGNORECASE),
        re.compile(r'(\d+|одного|один|двух|два|трех|три|четырех|четыре|пят[ьи]|шест[иь]|сем[иь]|восьм[иь]|девят[иь]|десят[иь])\s(дней|месяцев|лет)\sназад',
                   re.IGNORECASE),

    ]
    year_from_user_date = user_date.split('.')[-1]
    date_pattern = re.compile(r'^\d{2}\.\d{2}\.\d{4}')
    if date_pattern.match(paragraph.strip()):
        return paragraph

    for pattern in patterns:
        match = re.search(pattern, paragraph)
        if match:
            if pattern in [patterns[0], patterns[1]]:
                return f'{user_date} {paragraph}'
            elif pattern == patterns[2]:
                day = match.group(1)
                month_str = match.group(2).lower()
                year = match.group(3)
                month_num = months.get(month_str, 0)
                new_date = f'{int(day):02}.{month_num:02}.{year}'
                return f'{new_date} {paragraph}'
            elif pattern == patterns[3]:
                day1 = match.group(1)
                day2 = match.group(2)
                month_str = match.group(3).lower()
                month_num = months.get(month_str, 0)
                year = year_from_user_date
                new_date = f'{int(day1):02}.{month_num:02}.{year}'
                return f'{new_date} {paragraph}'
            elif pattern == patterns[4]:
                day, month, year = match.groups()
                new_date = f'{day}.{month}.{year}'
                return f'{new_date} {paragraph}'
            elif pattern == patterns[5]:
                amount_str = match.group(1).lower()
                amount = word_to_num.get(amount_str, 0)
                unit = match.group(2).lower()
                new_date = add_months_or_years_to_date(user_date, amount, unit)
                return f'{new_date} {paragraph}'
            elif pattern == patterns[6]:
                month1_str = match.group(1).lower()
                year = match.group(3)
                month1_num = months.get(month1_str, 0)
                new_date = f'01.{month1_num:02}.{year}'
                return f'{new_date} {paragraph}'
            elif pattern == patterns[7]:
                day, month, year = match.groups()
                # Преобразуем год в формат YYYY
                year = f'20{year}' if int(year) < 50 else f'19{year}'
                new_date = f'{day}.{month}.{year}'
                return f'{new_date} {paragraph}'
            elif pattern == patterns[8]:
                season, year = match.groups()
                month_num = seasons[season.lower()]
                new_date = f'01.{month_num:02}.{year}'
                return f'{new_date} {paragraph}'
            elif pattern == patterns[9]:
                day, month_str = match.groups()
                month_num = months.get(month_str.lower())
                new_date = f'{int(day):02}.{month_num:02}.{year_from_user_date}'
                return f'{new_date} {paragraph}'
            elif pattern == patterns[10]:
                month, year = match.groups()
                month_num = months[month.lower()]
                new_date = f'01.{month_num:02}.{year}'
                return f'{new_date} {paragraph}'
            elif pattern == patterns[11]:
                day, month, year = match.groups()
                # Преобразуем год в формат YYYY
                year = f'20{year}' if int(year) < 50 else f'19{year}'
                new_date = f'{day}.{month}.{year}'
                return f'{new_date} {paragraph}'
            elif pattern == patterns[12]:
                month_str = match.group(1).lower()
                month_num = months.get(month_str, 0)
                new_date = f'01.{month_num:02}.{int(year_from_user_date) - 1}'
                return f'{new_date} {paragraph}'
            elif pattern == patterns[13]:
                month_str = match.group(1).lower()
                month_num = months.get(month_str, 0)
                new_date = f'01.{month_num:02}.{int(year_from_user_date) + 1}'
                return f'{new_date} {paragraph}'
            elif pattern == patterns[14]:
                year = match.group(1)
                date_obj = parse(user_date, dayfirst=True)
                new_date_obj = date_obj - relativedelta(days=7)
                new_date = new_date_obj.strftime(f'%d.%m.{year}')
                return f'{new_date} {paragraph}'
            elif pattern == patterns[15]:
                day = match.group(1)
                month_str = match.group(2).lower()
                year = match.group(3)
                month_num = months.get(month_str, 0)
                new_date = f'{int(day):02}.{month_num:02}.{year}'
                return f'{new_date} {paragraph}'
            elif pattern == patterns[16]:
                day = match.group(1)
                month_str = match.group(2).lower()
                year_str = match.group(3).lower()
                month_num = months.get(month_str, 0)

                if year_str == 'этого года':
                    year = year_from_user_date
                elif year_str == 'предыдущего года':
                    year = year_from_user_date - 1
                elif year_str == 'следующего года':
                    year = year_from_user_date + 1

                new_date = f'{int(day):02}.{month_num:02}.{year}'
                return f'{new_date} {paragraph}'
            elif pattern == patterns[17]:
                season = match.group(1).lower()
                year_str = match.group(2).lower()
                month_num = seasons.get(season, 0)
                if year_str == 'этого года':
                    year = year_from_user_date
                elif year_str == 'следующего года':
                    year = year_from_user_date + 1
                elif year_str == 'предыдущего года':
                    year = year_from_user_date - 1

                new_date = f'01.{month_num:02}.{year}'
                return f'{new_date} {paragraph}'
            elif pattern == patterns[18]:
                day, month, year = match.groups()
                year = f'20{year}'
                new_date = f'{day}.{month}.{year}'
                return f'{new_date} {paragraph}'
            elif pattern == patterns[19]:
                year = match.group(1)
                new_date = f'01.01.{year}'
                return f'{new_date} {paragraph}'
            elif pattern == patterns[20]:
                phrase = match.group(0).lower()
                date_obj = parse(user_date, dayfirst=True)

                if phrase == 'вчера':
                    new_date_obj = date_obj - relativedelta(days=1)
                elif phrase == 'позавчера':
                    new_date_obj = date_obj - relativedelta(days=2)
                elif phrase == 'завтра':
                    new_date_obj = date_obj + relativedelta(days=1)
                elif phrase == 'послезавтра':
                    new_date_obj = date_obj + relativedelta(days=2)
                new_date = new_date_obj.strftime('%d.%m.%Y')
                return f'{new_date} {paragraph}'
            elif pattern == patterns[21]:
                number = match.group(1)
                unit = match.group(2).lower()
                date_obj = parse(user_date, dayfirst=True)

                if number is None:
                    number = 1

                number = int(number)

                if unit.startswith('месяц'):
                    new_date_obj = date_obj + relativedelta(months=number)
                elif unit in ['год', 'года', 'лет']:
                    new_date_obj = date_obj + relativedelta(years=number)

                new_date = new_date_obj.strftime('%d.%m.%Y')
                return f'{new_date} {paragraph}'
            elif pattern == patterns[22]:
                phrase = match.group(1).lower()
                date_obj = parse(user_date, dayfirst=True)

                if phrase == 'день':
                    number = 1
                elif phrase == 'несколько дней':
                    number = 2
                else:
                    number = int(re.match(r'(\д+)', phrase).group(1))

                new_date_obj = date_obj + relativedelta(days=number)
                new_date = new_date_obj.strftime('%d.%m.%Y')
                return f'{new_date} {paragraph}'
            elif pattern == patterns[23]:
                number_str = match.group(1).lower()
                unit = match.group(2).lower()
                date_obj = parse(user_date, dayfirst=True)

                if number_str.isdigit():
                    number = int(number_str)
                else:
                    number = word_to_num.get(number_str, 0)

                if unit == 'дней':
                    new_date_obj = date_obj - relativedelta(days=number)
                elif unit == 'месяцев':
                    new_date_obj = date_obj - relativedelta(months=number)
                elif unit == 'лет':
                    new_date_obj = date_obj - relativedelta(years=number)

                new_date = new_date_obj.strftime('%d.%m.%Y')
                return f'{new_date} {paragraph}'
    return paragraph


def duplicate_dates(input_docx_path, output_docx_path):
    doc = docx.Document(input_docx_path)
    return doc

def process_paragraph_with_month(paragraph, user_date):
    sentences = paragraph.split('.')
    processed_sentences = []
    date_pattern = re.compile(r'^\d{2}\.\d{2}\.\d{4}')
    if date_pattern.match(paragraph.strip()):
        return paragraph

    for sentence in sentences:
        match_first_half = re.search(r'на\s(январ|феврал|март|апрел|ма|июн|июл|август|сентябр|октябр|ноябр|декабр)', sentence, re.IGNORECASE)
        match_second_half = re.search(r'вторая половина\s(январ|феврал|март|апрел|ма|июн|июл|август|сентябр|октябр|ноябр|декабр)', sentence, re.IGNORECASE)

        if match_second_half:
            month_prefix = match_second_half.group(1)
            month_mapping = {
                'январ': '01', 'феврал': '02', 'март': '03', 'апрел': '04',
                'ма': '05', 'июн': '06', 'июл': '07', 'август': '08',
                'сентябр': '09', 'октябр': '10', 'ноябр': '11', 'декабр': '12'
            }
            month_number = month_mapping.get(month_prefix.lower())
            year_match = re.search(r'\d{4}', user_date)

            if month_number and year_match:
                year = year_match.group()
                converted_date = f'15.{month_number}.{year}'
                updated_sentence = f'{converted_date} {sentence}'
                processed_sentences.append(updated_sentence)
            else:
                processed_sentences.append(sentence)
        elif match_first_half:
            month_prefix = match_first_half.group(1)
            month_mapping = {
                'январ': '01', 'феврал': '02', 'март': '03', 'апрел': '04',
                'ма': '05', 'июн': '06', 'июл': '07', 'август': '08',
                'сентябр': '09', 'октябр': '10', 'ноябр': '11', 'декабр': '12'
            }
            month_number = month_mapping.get(month_prefix.lower())
            year_match = re.search(r'\d{4}', user_date)

            if month_number and year_match:
                year = year_match.group()
                converted_date = f'01.{month_number}.{year}'
                updated_sentence = f'{converted_date} {sentence}'
                processed_sentences.append(updated_sentence)
            else:
                processed_sentences.append(sentence)
        else:
            processed_sentences.append(sentence)

    return '.'.join(processed_sentences).strip()


from dateutil.parser import parse
from dateutil.relativedelta import relativedelta


def add_months_or_years_to_date(date_str, amount, unit):
    date_obj = parse(date_str, dayfirst=True)
    if unit == 'месяц' or unit == 'месяца' or unit == 'месяцев':
        new_date_obj = date_obj + relativedelta(months=amount)
    elif unit == 'год' or unit == 'года' or unit == 'лет':
        new_date_obj = date_obj + relativedelta(years=amount)
    return new_date_obj.strftime('%d.%m.%Y')

user_date = '31.03.2024'
input_docx_path = "result.docx"
output_docx_path = "result_with_dates.docx"

doc = duplicate_dates(input_docx_path, output_docx_path)
result_doc = docx.Document()

for para in doc.paragraphs:
    new_para = result_doc.add_paragraph()
    updated_para_text = process_today_date_in_text(para.text, user_date)
    para.text = updated_para_text
    if updated_para_text != para.text:
        new_para.add_run(updated_para_text)
    new_para.add_run(process_paragraph_with_month(para.text, user_date))


def extract_date(paragraph_text):
    match = re.match(r'(\d{2}\.\d{2}\.\d{4})', paragraph_text)
    if match:
        return parse(match.group(1), dayfirst=True)
    return None

def sort_paragraphs_by_date(paragraphs):
    dated_paragraphs = [(extract_date(para.text), para.text) for para in paragraphs if extract_date(para.text)]
    sorted_paragraphs = sorted(dated_paragraphs, key=lambda x: x[0])
    return [para for date, para in sorted_paragraphs]

# Извлечение всех параграфов
paragraphs = result_doc.paragraphs

# Сортировка параграфов по дате
sorted_paragraphs = sort_paragraphs_by_date(paragraphs)

# Создание нового документа и добавление отсортированных параграфов
sorted_doc = Document()
for para_text in sorted_paragraphs:
    sorted_doc.add_paragraph(para_text)

# Сохранение отсортированного документа
sorted_doc.save('sorted_document.docx')

result_doc.save("result_with_dates.docx")